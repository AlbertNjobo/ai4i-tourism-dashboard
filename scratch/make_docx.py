import re
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>\n'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>\n'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>\n'
        f'  <w:left w:val="none"/>\n'
        f'  <w:right w:val="none"/>\n'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="EAEAEA"/>\n'
        f'  <w:insideV w:val="none"/>\n'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def build_proposal_docx(md_path, docx_path):
    doc = Document()
    
    # 1. Page Setup: 1 inch margins all sides
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11.0)

    # 2. Typography: Arial, 11pt, 1.15 line spacing
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Arial'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    style_normal.paragraph_format.line_spacing = 1.15
    style_normal.paragraph_format.space_after = Pt(8)

    # Load content
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    # Split lines and parse
    lines = content.split('\n')
    
    in_mermaid = False
    in_table = False
    table_headers = []
    table_rows = []
    
    idx = 0
    while idx < len(lines):
        line = lines[idx].strip()
        
        # Handle mermaid blocks
        if line.startswith('```mermaid'):
            in_mermaid = True
            idx += 1
            mermaid_lines = []
            while idx < len(lines) and not lines[idx].strip().startswith('```'):
                mermaid_lines.append(lines[idx])
                idx += 1
            # Add mermaid block as custom italic block
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run("[System Flow Diagram]\n" + "\n".join(mermaid_lines))
            run.font.italic = True
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            in_mermaid = False
            idx += 1
            continue

        # Skip other code blocks (just plain text for this doc)
        if line.startswith('```'):
            idx += 1
            code_lines = []
            while idx < len(lines) and not lines[idx].strip().startswith('```'):
                code_lines.append(lines[idx])
                idx += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            run = p.add_run("\n".join(code_lines))
            run.font.name = 'Courier New'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            idx += 1
            continue

        # Check for table
        if line.startswith('|') and idx < len(lines):
            in_table = True
            # Read headers
            header_line = line
            sep_line = lines[idx+1].strip()
            
            # Parse header columns
            headers = [c.strip() for c in header_line.split('|')[1:-1]]
            
            # Read data rows
            rows = []
            row_idx = idx + 2
            while row_idx < len(lines) and lines[row_idx].strip().startswith('|'):
                row_cols = [c.strip() for c in lines[row_idx].split('|')[1:-1]]
                rows.append(row_cols)
                row_idx += 1
            
            # Generate docx table
            doc_table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
            doc_table.autofit = True
            set_table_borders(doc_table)
            
            # Set Header Row
            hdr_cells = doc_table.rows[0].cells
            for col_i, h_text in enumerate(headers):
                hdr_cells[col_i].text = h_text
                set_cell_background(hdr_cells[col_i], 'F1F5F9')
                set_cell_margins(hdr_cells[col_i], top=120, bottom=120, left=150, right=150)
                # Bold headers
                for paragraph in hdr_cells[col_i].paragraphs:
                    paragraph.paragraph_format.space_after = Pt(0)
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(10)
                        run.font.name = 'Arial'

            # Set Data Rows
            for r_i, r_cols in enumerate(rows):
                row_cells = doc_table.rows[r_i + 1].cells
                bg_color = 'FFFFFF' if r_i % 2 == 0 else 'F8FAFC'
                for col_i, cell_val in enumerate(r_cols):
                    if col_i < len(row_cells):
                        row_cells[col_i].text = cell_val
                        set_cell_background(row_cells[col_i], bg_color)
                        set_cell_margins(row_cells[col_i], top=100, bottom=100, left=150, right=150)
                        for paragraph in row_cells[col_i].paragraphs:
                            paragraph.paragraph_format.space_after = Pt(0)
                            for run in paragraph.runs:
                                run.font.size = Pt(9.5)
                                run.font.name = 'Arial'
            
            # Add spacer after table
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_before = Pt(8)
            spacer.paragraph_format.space_after = Pt(8)

            idx = row_idx
            in_table = False
            continue

        # Headings
        if line.startswith('# '):
            h_text = line[2:]
            p = doc.add_heading(level=1)
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(h_text)
            run.font.name = 'Arial'
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A) # Slate 900
            
            # If it is the main title, center it and add a page break after cover info
            if h_text == 'Zimbabwe Tourism Destination Insights':
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Center cover page items
                idx += 1
                cover_lines = []
                while idx < len(lines) and not lines[idx].strip().startswith('---'):
                    line_sub = lines[idx].strip()
                    if line_sub:
                        cover_lines.append(line_sub)
                    idx += 1
                
                # Add centered subheadings for cover page
                for c_line in cover_lines:
                    c_line_cleaned = c_line.replace('**', '').strip()
                    cp = doc.add_paragraph()
                    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cp.paragraph_format.space_after = Pt(6)
                    run_c = cp.add_run(c_line_cleaned)
                    run_c.font.size = Pt(12)
                    run_c.font.name = 'Arial'
                    if ':' in c_line_cleaned:
                        run_c.font.bold = True
                
                # Add Page Break to start Section 1 on page 2
                doc.add_page_break()
                idx += 1
                continue
                
        elif line.startswith('## '):
            h_text = line[3:]
            p = doc.add_heading(level=2)
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(h_text)
            run.font.name = 'Arial'
            run.font.size = Pt(13)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B) # Slate 800
        elif line.startswith('### '):
            h_text = line[4:]
            p = doc.add_heading(level=3)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(h_text)
            run.font.name = 'Arial'
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x33, 0x41, 0x55) # Slate 700

        # Bullet lists
        elif line.startswith('- '):
            # Clean markdown bold/links
            cleaned = line[2:]
            cleaned = cleaned.replace('**', '')
            # Clean links [label](url) -> label
            cleaned = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', cleaned)
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(cleaned)
            run.font.name = 'Arial'
            run.font.size = Pt(11)

        # Standard paragraph
        elif line:
            # Skip horizontal rules
            if line == '---':
                idx += 1
                continue
            
            cleaned = line.replace('**', '')
            cleaned = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', cleaned)
            
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run(cleaned)
            run.font.name = 'Arial'
            run.font.size = Pt(11)
            
        idx += 1

    # Save docx
    os.makedirs(os.path.dirname(docx_path), exist_ok=True)
    doc.save(docx_path)
    print(f"Document successfully created at: {docx_path}")

if __name__ == '__main__':
    md_file = '/home/lawrence/Downloads/AI4I ToRs-20260710T233114Z-2-001/AI4I ToRs/ai4i-tourism-dashboard/docs/proposal_draft.md'
    docx_file = '/home/lawrence/Downloads/AI4I ToRs-20260710T233114Z-2-001/AI4I ToRs/ai4i-tourism-dashboard/docs/[ProjectID]_AI4I_Proposal_Design.docx'
    build_proposal_docx(md_file, docx_file)
